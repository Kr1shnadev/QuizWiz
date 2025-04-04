import os
import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

# Define save directory
SAVE_DIR = os.path.join(os.path.expanduser('~'), 'Downloads')
os.makedirs(SAVE_DIR, exist_ok=True)

def generate_pdf(sections, blueprint, filename):
    # Create PDF file path
    filename = f"{blueprint.get('examName', 'question_paper')}.pdf"
    filepath = os.path.join(SAVE_DIR, filename)
    
    # Create the PDF document
    doc = SimpleDocTemplate(filepath, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    
    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Title'],
        fontSize=16,
        spaceAfter=30
    )
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading1'],
        fontSize=14,
        spaceAfter=20
    )
    normal_style = styles['Normal']
    
    # Add institution name
    story.append(Paragraph('QuizWiz', title_style))
    
    # Add header
    story.append(Paragraph('Internal Exam', heading_style))
    
    # Add exam details
    details_text = f"Time: {blueprint['duration']} | Max Marks: {blueprint['totalMarks']}"
    story.append(Paragraph(details_text, normal_style))
    
    # Add date
    date_text = f"Date: {datetime.datetime.now().strftime('%d-%m-%Y')}"
    story.append(Paragraph(date_text, normal_style))
    story.append(Spacer(1, 0.5*inch))
    
    # Add sections and questions
    question_number = 1
    for section in sections:
        # Add section header
        story.append(Paragraph(f"Section {section['sectionName']}", heading_style))
        
        # Add questions
        for question in section['questions']:
            question_text = f"{question_number}. {question['text']} ({question['marks']} marks)"
            story.append(Paragraph(question_text, normal_style))
            
            # Add options if multiple choice
            if 'options' in question:
                for option in question['options']:
                    story.append(Paragraph(f"    {option}", normal_style))
            
            story.append(Spacer(1, 0.2*inch))
            question_number += 1
    
    # Build the PDF
    doc.build(story)
    return filepath

def generate_docx(sections, blueprint, filename):
    doc = Document()
    
    # Add institution name
    institution = doc.add_paragraph()
    institution.alignment = WD_ALIGN_PARAGRAPH.CENTER
    institution_run = institution.add_run('QuizWiz')
    institution_run.bold = True
    institution_run.font.size = Pt(16)
    
    # Add header
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header.add_run('Internal Exam')
    header_run.bold = True
    header_run.font.size = Pt(14)
    
    # Add exam details
    details = doc.add_paragraph()
    details.alignment = WD_ALIGN_PARAGRAPH.CENTER
    details_text = f"Time: {blueprint['duration']} | Max Marks: {blueprint['totalMarks']}"
    details_run = details.add_run(details_text)
    details_run.font.size = Pt(12)
    
    # Add date
    date = doc.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_text = f"Date: {datetime.datetime.now().strftime('%d-%m-%Y')}"
    date_run = date.add_run(date_text)
    date_run.font.size = Pt(12)
    
    doc.add_paragraph()  # Add spacing
    
    # Add sections and questions
    question_number = 1
    for section in sections:
        # Add section header
        section_para = doc.add_paragraph()
        section_run = section_para.add_run(f"Section {section['sectionName']}")
        section_run.bold = True
        
        # Add questions
        for question in section['questions']:
            question_para = doc.add_paragraph()
            question_para.add_run(f"{question_number}. {question['text']} ")
            question_para.add_run(f"({question['marks']} marks)")
            
            # Add options if multiple choice
            if 'options' in question:
                for option in question['options']:
                    option_para = doc.add_paragraph()
                    option_para.paragraph_format.left_indent = Inches(0.5)
                    option_para.add_run(option)
            
            question_number += 1
    
    # Save document with custom filename
    filename = f"{blueprint.get('examName', 'question_paper')}.docx"
    filepath = os.path.join(SAVE_DIR, filename)
    doc.save(filepath)
    return filepath



# Main entry point for command line usage
if __name__ == '__main__':
    import sys
    import json
    
    # Parse the JSON input from command line
    if len(sys.argv) > 1:
        try:
            data = json.loads(sys.argv[1])
            sections = data['sections']
            blueprint = data['blueprint']
            format_type = data['format']
            filename = data['filename']
            
            if format_type == 'pdf':
                filepath = generate_pdf(sections, blueprint, filename)
            else:
                filepath = generate_docx(sections, blueprint, filename)
                
            print(filepath)  # Return the filepath to JavaScript
        except Exception as e:
            print(f'Error: {str(e)}')
            sys.exit(1)

# Export functions for use in JavaScript
class QuestionPaperGenerator:
    @staticmethod
    def generateDOCX(sections, blueprint, filename):
        return generate_docx(sections, blueprint, filename)
        
    @staticmethod
    def generatePDF(sections, blueprint, filename):
        return generate_pdf(sections, blueprint, filename)